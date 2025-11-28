import re
import io
import requests
import streamlit as st
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from PIL import Image
import matplotlib.pyplot as plt

st.set_page_config(page_title="Sinh Đề GDCD Tự Động", page_icon="📚", layout="wide")
st.title("📚 Sinh Đề GDCD – LaTeX → ảnh → DOCX/PDF")

# --- API KEY ---
api_key = st.secrets.get("GOOGLE_API_KEY", "")
if not api_key:
    api_key = st.text_input("Nhập Google API Key:", type="password")

# --- GUI ---
lop_options = ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"]

chuong_options = {
    "Lớp 6": [
        "Chủ đề 1: Quyền và nghĩa vụ cơ bản của công dân",
        "Chủ đề 2: Kỷ luật, pháp luật và trách nhiệm",
        "Chủ đề 3: Đạo đức trong học tập và đời sống"
    ],
    "Lớp 7": [
        "Chủ đề 1: Quyền và nghĩa vụ trong trường học",
        "Chủ đề 2: Kỹ năng sống cơ bản",
        "Chủ đề 3: Xây dựng môi trường văn hóa"
    ],
    "Lớp 8": [
        "Chủ đề 1: Công dân và pháp luật",
        "Chủ đề 2: Đạo đức nghề nghiệp và trách nhiệm xã hội",
        "Chủ đề 3: An toàn và bảo vệ môi trường"
    ],
    "Lớp 9": [
        "Chủ đề 1: Quyền và nghĩa vụ công dân trong xã hội",
        "Chủ đề 2: Pháp luật và hình thức xử lý vi phạm",
        "Chủ đề 3: Xây dựng nếp sống văn minh"
    ]
}

bai_options = {
    # --- Lớp 6 ---
    "Chủ đề 1: Quyền và nghĩa vụ cơ bản của công dân": ["Bài 1: Quyền cơ bản", "Bài 2: Nghĩa vụ cơ bản"],
    "Chủ đề 2: Kỷ luật, pháp luật và trách nhiệm": ["Bài 1: Kỷ luật ở trường học", "Bài 2: Pháp luật cơ bản"],
    "Chủ đề 3: Đạo đức trong học tập và đời sống": ["Bài 1: Trung thực và tôn trọng", "Bài 2: Giúp đỡ bạn bè"],

    # --- Lớp 7 ---
    "Chủ đề 1: Quyền và nghĩa vụ trong trường học": ["Bài 1: Quyền học tập", "Bài 2: Nghĩa vụ học tập"],
    "Chủ đề 2: Kỹ năng sống cơ bản": ["Bài 1: Giao tiếp", "Bài 2: Giải quyết mâu thuẫn"],
    "Chủ đề 3: Xây dựng môi trường văn hóa": ["Bài 1: Văn hóa học đường", "Bài 2: Hoạt động tập thể"],

    # --- Lớp 8 ---
    "Chủ đề 1: Công dân và pháp luật": ["Bài 1: Luật pháp cơ bản", "Bài 2: Trách nhiệm tuân thủ pháp luật"],
    "Chủ đề 2: Đạo đức nghề nghiệp và trách nhiệm xã hội": ["Bài 1: Đạo đức nghề nghiệp", "Bài 2: Trách nhiệm xã hội"],
    "Chủ đề 3: An toàn và bảo vệ môi trường": ["Bài 1: An toàn cá nhân", "Bài 2: Bảo vệ môi trường"],

    # --- Lớp 9 ---
    "Chủ đề 1: Quyền và nghĩa vụ công dân trong xã hội": ["Bài 1: Quyền công dân", "Bài 2: Nghĩa vụ công dân"],
    "Chủ đề 2: Pháp luật và hình thức xử lý vi phạm": ["Bài 1: Hình thức xử lý", "Bài 2: Trách nhiệm pháp lý"],
    "Chủ đề 3: Xây dựng nếp sống văn minh": ["Bài 1: Văn minh nơi công cộng", "Bài 2: Nếp sống văn hóa"]
}

with st.sidebar:
    st.header("Thông tin sinh đề")
    lop = st.selectbox("Chọn lớp", lop_options)
    chuong = st.selectbox("Chọn chủ đề/chương", chuong_options[lop])
    bai_list = bai_options.get(chuong, [])
    if bai_list:
        bai = st.selectbox("Chọn bài", bai_list)
    else:
        bai = st.text_input("Chưa có bài cho chủ đề này", "")

    so_cau = st.number_input("Số câu hỏi", min_value=1, max_value=50, value=10)
    loai_cau = st.selectbox(
        "Loại câu hỏi",
        ["Trắc nghiệm 4 lựa chọn", "Trắc nghiệm Đúng – Sai", "Câu trả lời ngắn", "Tự luận", "Trộn ngẫu nhiên"]
    )
    co_dap_an = st.checkbox("Có đáp án", value=True)

# --- BUILD PROMPT ---
def build_prompt(lop, chuong, bai, so_cau, loai_cau, co_dap_an):
    return f"""
Bạn là giáo viên GDCD. Hãy sinh đề kiểm tra:
- Lớp: {lop}
- Chủ đề/Chương: {chuong}
- Bài: {bai}
- Số câu hỏi: {so_cau}
- Loại câu hỏi: {loai_cau}
- {"Có đáp án" if co_dap_an else "Không có đáp án"}

YÊU CẦU QUAN TRỌNG:
1) Toàn bộ công thức (nếu có) phải viết bằng LaTeX $$...$$.
2) Câu trắc nghiệm: A. ... B. ... C. ... D. ...
3) Câu trả lời ngắn: 1 dòng.
4) Đáp án dưới câu hỏi, cách 2 dòng trống.
5) Chỉ dùng tiếng Việt.
"""

# --- Gọi API ---
def generate_questions(api_key, lop, chuong, bai, so_cau, loai_cau, co_dap_an):
    MODEL = "models/gemini-2.0-flash"
    url = f"https://generativelanguage.googleapis.com/v1/{MODEL}:generateContent?key={api_key}"
    prompt = build_prompt(lop, chuong, bai, so_cau, loai_cau, co_dap_an)
    payload = {"contents":[{"role":"user","parts":[{"text":prompt}]}]}
    try:
        r = requests.post(url, json=payload, timeout=30)
        if r.status_code != 200:
            return f"❌ Lỗi API {r.status_code}: {r.text}"
        j = r.json()
        return j["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        return f"❌ Lỗi kết nối: {e}"

# --- Xử lý LaTeX ---
LATEX_RE = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)
def find_latex_blocks(text):
    return [(m.span(), m.group(0), m.group(1)) for m in LATEX_RE.finditer(text)]

def render_latex_png_bytes(latex_code, fontsize=20, dpi=200):
    fig = plt.figure()
    fig.patch.set_alpha(0.0)
    fig.text(0, 0, f"${latex_code}$", fontsize=fontsize)
    buf = io.BytesIO()
    plt.axis('off')
    plt.savefig(buf, format='png', dpi=dpi, bbox_inches='tight', pad_inches=0.02, transparent=True)
    plt.close(fig)
    buf.seek(0)
    return buf.read()

def create_docx_bytes(text):
    doc = Document()
    last = 0
    for span, full, inner in find_latex_blocks(text):
        start, end = span
        before = text[last:start]
        for line in before.splitlines():
            doc.add_paragraph(line)
        try:
            png_bytes = render_latex_png_bytes(inner)
            img_stream = io.BytesIO(png_bytes)
            p = doc.add_paragraph()
            r = p.add_run()
            r.add_picture(img_stream, width=Inches(3))
        except:
            doc.add_paragraph(full)
        last = end
    for line in text[last:].splitlines():
        doc.add_paragraph(line)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

def create_pdf_bytes(text):
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    margin = 40
    y = height - 50
    last = 0
    for span, full, inner in find_latex_blocks(text):
        start, end = span
        before = text[last:start]
        for line in before.splitlines():
            c.drawString(margin, y, line)
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 50
        try:
            png_bytes = render_latex_png_bytes(inner)
            img_reader = ImageReader(io.BytesIO(png_bytes))
            img = Image.open(io.BytesIO(png_bytes))
            draw_w = 300
            draw_h = img.height / img.width * draw_w
            if y - draw_h < 60:
                c.showPage()
                y = height - 50
            c.drawImage(img_reader, margin, y - draw_h, width=draw_w, height=draw_h, mask='auto')
            y -= draw_h + 8
        except:
            c.drawString(margin, y, full)
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 50
        last = end
    for line in text[last:].splitlines():
        c.drawString(margin, y, line)
        y -= 14
        if y < 60:
            c.showPage()
            y = height - 50
    c.save()
    buf.seek(0)
    return buf

# --- BUTTON ---
if st.button("🎯 Sinh đề ngay"):
    if not api_key:
        st.error("Thiếu API Key!")
    else:
        with st.spinner("⏳ AI đang tạo đề..."):
            result = generate_questions(api_key, lop, chuong, bai, so_cau, loai_cau, co_dap_an)

        if isinstance(result, str) and result.startswith("❌"):
            st.error(result)
        else:
            st.success("🎉 Đã tạo xong đề (hiển thị nội dung).")
            st.markdown(result.replace("\n", "<br>"), unsafe_allow_html=True)

            latex_blocks = find_latex_blocks(result)
            if not latex_blocks:
                st.warning("Không tìm thấy LaTeX. Xuất TXT.")
                st.download_button(
                    "📥 Tải TXT", data=result.encode("utf-8"),
                    file_name=f"De_{lop}_{chuong}_{bai}.txt", mime="text/plain"
                )
            else:
                try:
                    docx_io = create_docx_bytes(result)
                    st.download_button(
                        "📥 Tải DOCX",
                        data=docx_io.getvalue(),
                        file_name=f"De_{lop}_{chuong}_{bai}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Không tạo DOCX: {e}")

                try:
                    pdf_io = create_pdf_bytes(result)
                    st.download_button(
                        "📥 Tải PDF",
                        data=pdf_io.getvalue(),
                        file_name=f"De_{lop}_{chuong}_{bai}.pdf",
                        mime="application/pdf"
                    )
                except Exception as e:
                    st.error(f"Không tạo PDF: {e}")
